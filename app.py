import streamlit as st
import google.generativeai as genai
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, Inches, RGBColor
import re
import base64
import tempfile
import os

# ====================== CẤU HÌNH ======================
st.set_page_config(page_title="Soạn Đề Siêu Dễ - THCS Tân Hội Đông", page_icon="test_tube", layout="wide")

st.markdown("""
<style>
    .main-title {font-size: 3.2rem; color: #c62828; text-align: center; font-weight: bold; margin: 30px 0;}
    .sub-title {font-size: 1.6rem; color: #1565c0; text-align: center; font-weight: bold; margin-bottom: 40px;}
    .stButton>button {background: linear-gradient(45deg, #c62828, #e53935); color: white; font-weight: bold; height: 70px; font-size: 22px; border-radius: 15px;}
    .preview-box {background-color: #fff8f8; padding: 25px; border-radius: 12px; border-left: 8px solid #c62828; margin: 20px 0;}
    .correct {color: red !important; font-weight: bold !important;}
    .question {font-weight: bold; color: #1976d2; font-size: 1.2em; margin: 15px 0 8px 0;}
    .option {margin: 6px 0;}
    .info-box {background-color: #e3f2fd; padding: 15px; border-radius: 10px; border-left: 5px solid #1976d2; margin: 15px 0;}
</style>
""", unsafe_allow_html=True)

# ====================== API KEY ======================
if "GEMINI_API_KEY" not in st.secrets:
    st.error("Lỗi: Chưa có GEMINI_API_KEY trong Secrets!")
    st.info("Vào **Settings → Secrets** → thêm dòng:\nGEMINI_API_KEY = \"AIzaSy...\"")
    st.stop()
genai.configure(api_key=st.secrets["GEMINI_API_KEY"])

# ====================== ĐỌC FILE ======================
def read_file(file):
    if file.name.endswith(".pdf"):
        try: return " ".join([p.extract_text() or "" for p in PdfReader(file).pages])
        except: return None
    else:
        try: return "\n".join([p.text for p in Document(file).paragraphs if p.text.strip()])
        except: return None

# ====================== TIÊU ĐỀ ======================
st.markdown('<div class="main-title">SOẠN ĐỀ TRẮC NGHIỆM</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-title">Bản Demo của Trường THCS Tân Hội Đông - Đồng Tháp</div>', unsafe_allow_html=True)

# ====================== GIAO DIỆN ======================
col1, col2 = st.columns([1, 1])

with col1:
    st.subheader("1. Nhập nội dung")
    objectives = st.text_area("Yêu cầu cần đạt (bắt buộc):", height=160)
    uploaded = st.file_uploader("Tải giáo án (PDF/Word)", type=["pdf", "docx"])
    content = read_file(uploaded) if uploaded else ""
    if uploaded and content:
        st.success(f"Đã đọc: {uploaded.name}")

with col2:
    st.subheader("2. Chỉ cần TÍCH vào ô vuông")

    # BỔ SUNG ĐỊNH NGHĨA 3 MỨC ĐỘ
    with st.expander("Nhấp để xem định nghĩa 3 mức độ nhận thức", expanded=False):
        st.markdown("""
        <div class="info-box">
        <strong>1. Biết (Nhận biết)</strong><br>
        • Mức độ: Thấp nhất<br>
        • Yêu cầu: Ghi nhớ, nhận diện, phát biểu lại các khái niệm, sự kiện, công thức đã học<br>
        • Từ khóa: Định nghĩa, mô tả, liệt kê, gọi tên, chọn ra<br>
        • Ví dụ: "Thủ đô của Việt Nam là gì?" → Hà Nội
        
        <strong>2. Hiểu (Thông hiểu)</strong><br>
        • Mức độ: Cao hơn Biết<br>
        • Yêu cầu: Diễn giải, tóm tắt, giải thích được ý nghĩa của thông tin đã biết<br>
        • Từ khóa: Diễn giải, tóm tắt, giải thích, so sánh<br>
        • Ví dụ: "Giải thích tại sao Hà Nội lại là thủ đô của Việt Nam"
        
        <strong>3. Vận dụng</strong><br>
        • Mức độ: Cao nhất<br>
        • Yêu cầu: Sử dụng kiến thức đã học để giải quyết vấn đề mới<br>
        • Phân loại:<br>
          &nbsp;&nbsp;→ Vận dụng đơn giản: Áp dụng vào tình huống quen thuộc<br>
          &nbsp;&nbsp;→ Vận dụng nâng cao: Giải quyết vấn đề phức tạp, chưa từng gặp
        </div>
        """, unsafe_allow_html=True)

    st.markdown("#### Mức độ nhận thức:")
    lv1 = st.checkbox("**Biết** (Nhận biết – nhớ, gọi tên)", True)
    lv2 = st.checkbox("**Hiểu** (Thông hiểu – giải thích, so sánh)", True)
    lv3 = st.checkbox("**Vận dụng** (Giải quyết vấn đề mới)", False)
    selected_levels = [x for x, c in zip(["Biết","Hiểu","Vận dụng"], [lv1,lv2,lv3]) if c]
    st.success(f"Đã chọn: **{', '.join(selected_levels)}**")

    st.markdown("---")
    st.markdown("#### Loại câu hỏi:")
    t1 = st.checkbox("**4 đáp án (1 đúng)**", True)
    t2 = st.checkbox("**Đúng - Sai**", False)
    t3 = st.checkbox("**Nhiều lựa chọn đúng**", False)
    selected_types = []
    if t1: selected_types.append("4 đáp án")
    if t2: selected_types.append("Đúng/Sai")
    if t3: selected_types.append("Nhiều lựa chọn đúng")
    st.success(f"Đã chọn: **{', '.join(selected_types)}**")

    num_questions = st.slider("Số lượng câu hỏi:", 5, 25, 12)
    add_image = st.checkbox("Tự động thêm hình minh hoạ cho câu hình học", True)

# ====================== TẠO ĐỀ ======================
if st.button("TẠO ĐỀ & XUẤT FILE WORD NGAY", use_container_width=True):
    if not objectives.strip() or not content or not selected_levels or not selected_types:
        st.error("Vui lòng nhập đủ thông tin và chọn mức độ + loại câu hỏi!")
    else:
        with st.spinner("AI đang soạn đề chất lượng cao..."):
            model = genai.GenerativeModel("gemini-2.5-flash")
            prompt = f"""
            Soạn đúng {num_questions} câu trắc nghiệm, KHÔNG chào hỏi, KHÔNG giải thích.
            Bám sát: {objectives}
            Nội dung: {content[:25000]}

            Mức độ: {', '.join(selected_levels)}
            Loại: {', '.join(selected_types)}

            - Nếu là hình học → thêm dòng: [HÌNH MINH HỌA]
            - Công thức viết rõ: x², √2, △ABC, ∠A=90°
            - Mỗi đáp án 1 dòng: A. ..., B. ..., C. ..., D. ...
            - Ở cuối mỗi câu ghi đáp án đúng bằng ký hiệu: @@ĐÁP ÁN: C@@ (chỉ ghi đúng 1 chữ cái hoặc Đúng/Sai)
            """
            try:
                resp = model.generate_content(prompt)
                raw = resp.text
                st.session_state.raw = raw
                st.success("ĐÃ SOẠN XONG!")
            except Exception as e:
                st.error(f"Lỗi AI: {e}")
                st.stop()

        # ====================== XỬ LÝ ĐÁP ÁN ======================
        lines = st.session_state.raw.split('\n')
        questions = []
        current_q = []
        current_ans = None

        for line in lines:
            line = line.strip()
            if line.startswith("@@ĐÁP ÁN:"):
                current_ans = line.split("@@ĐÁP ÁN:")[1].split("@@")[0].strip().upper()
                if current_ans == "ĐÚNG": current_ans = "Đúng"
                if current_ans == "SAI": current_ans = "Sai"
                continue
            if line == "---" or not line:
                if current_q:
                    questions.append((current_q, current_ans))
                    current_q = []
                    current_ans = None
            else:
                current_q.append(line)

        if current_q:
            questions.append((current_q, current_ans))

        # ====================== XEM TRƯỚC ĐẸP ======================
        st.markdown("### XEM TRƯỚC ĐỀ (ĐÁP ÁN ĐÚNG ĐÃ TÔ ĐỎ)")
        st.markdown("<div class='preview-box'>", unsafe_allow_html=True)

        for q_lines, correct in questions:
            for line in q_lines:
                if line.lower().startswith("câu"):
                    st.markdown(f"<div class='question'>{line}</div>", unsafe_allow_html=True)
                elif "[hình minh họa" in line.lower():
                    st.markdown("**Hình minh hoạ sẽ được chèn trong file Word**")
                elif re.match(r'^[A-E]\.', line):
                    if correct and re.match(rf'^{re.escape(correct)}\.', line, re.IGNORECASE):
                        st.markdown(f"<div class='correct'>{line}</div>", unsafe_allow_html=True)
                    else:
                        st.markdown(f"<div class='option'>{line}</div>", unsafe_allow_html=True)
                else:
                    st.write(line)
            st.markdown("---")

        st.markdown("</div>", unsafe_allow_html=True)

        # ====================== TẠO FILE WORD ======================
        doc = Document()
        for s in doc.sections:
            s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(0.8)

        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(13)

        for q_lines, correct in questions:
            for line in q_lines:
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.space_after = Pt(6)

                if line.lower().startswith("câu"):
                    run = p.add_run(line)
                    run.bold = True

                elif "[hình minh họa" in line.lower() and add_image:
                    try:
                        img_model = genai.GenerativeModel("gemini-2.5-flash")
                        img = img_model.generate_content(
                            [f"Vẽ hình minh hoạ rõ nét cho: {line.replace('[HÌNH MINH HỌA]', '')[:200]}"],
                            generation_config={"response_mime_type": "image/jpeg"}
                        )
                        img_data = img.candidates[0].content.parts[0].inline_data.data
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as f:
                            f.write(base64.b64decode(img_data))
                            doc.add_picture(f.name, width=Inches(4))
                            os.unlink(f.name)
                        doc.add_paragraph()
                    except:
                        p.add_run(" (Hình vẽ tay nếu cần)")

                elif re.match(r'^[A-E]\.', line):
                    run = p.add_run(line)
                    if correct and re.match(rf'^{re.escape(correct)}\.', line, re.IGNORECASE):
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 0, 0)
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.style = 'List Bullet'

        # Lưu file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            path = tmp.name
        with open(path, "rb") as f:
            docx_bytes = f.read()
        os.unlink(path)

        # Nút tải
        st.download_button(
            label="TẢI FILE WORD (.docx) - Đáp án đúng đã tô đỏ trong A/B/C/D",
            data=docx_bytes,
            file_name="De_Trac_Nghiem_Tan_Hoi_Dong.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

        st.success("XONG 100%! Đáp án đúng đã tô đỏ – In thoải mái!")
        st.balloons()